import os
import shutil
import uuid
import tempfile
import re
from pathlib import Path
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional, List
import httpx
from bs4 import BeautifulSoup

# Import Gemini SDK
import google.generativeai as genai
from google.generativeai.client import FileServiceClient
import googleapiclient.http
import googleapiclient.discovery
import httplib2
from typing import Sequence

# Monkeypatch to support new Google "AQ." format API keys with the legacy SDK
def patched_setup_discovery_api(self, metadata: Sequence[tuple[str, str]] = ()):
    api_key = self._client_options.api_key
    if api_key is None:
        raise ValueError(
            "Invalid operation: Uploading to the File API requires an API key. Please provide a valid API key."
        )

    # Omit key query parameter when fetching the discovery doc because it causes 400 for new "AQ." keys
    discovery_url = "https://generativelanguage.googleapis.com/$discovery/rest?version=v1beta"
    request = googleapiclient.http.HttpRequest(
        http=httplib2.Http(),
        postproc=lambda resp, content: (resp, content),
        uri=discovery_url,
        headers=dict(metadata),
    )
    response, content = request.execute()
    request.http.close()

    discovery_doc = content.decode("utf-8")
    self._local.discovery_api = googleapiclient.discovery.build_from_document(
        discovery_doc, developerKey=api_key
    )

FileServiceClient._setup_discovery_api = patched_setup_discovery_api

# Load environment variables from .env.local
from dotenv import load_dotenv
env_path = Path(__file__).parent.parent / '.env.local'
load_dotenv(dotenv_path=env_path)

# Setup Gemini
# The API key must be set in the environment variable GEMINI_API_KEY
API_KEY = os.environ.get("GEMINI_API_KEY", "")
if API_KEY:
    genai.configure(api_key=API_KEY)

# Import prompts
from prompts import build_generation_prompt, build_refinement_prompt

def clean_ai_html_output(text: str) -> str:
    """
    Strip markdown backticks, conversational preamble, and enforce
    mandatory term replacements that the AI may ignore from prompts.
    """
    text = text.strip()

    # 1. Strip markdown code blocks (```html ... ```)
    match = re.search(r"```(?:html)?\s*(.*?)\s*```", text, re.DOTALL | re.IGNORECASE)
    if match and match.group(1).strip():
        text = match.group(1).strip()
    else:
        # Strip conversational text before/after the HTML
        start_idx = text.find("<")
        end_idx = text.rfind(">")
        if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
            stripped = text[start_idx:end_idx+1].strip()
            if stripped:
                text = stripped

    # 2. FORCE replace forbidden terms — normal case globally
    replacements = [
        (r"[Cc]ông an huyện\s*[^<;,\.\)\"]*", "Công an tỉnh Đắk Lắk"),
        (r"CÔNG AN HUYỆN\s*[^<;,\.\)\"]*", "Công an tỉnh Đắk Lắk"),
        (r"[Cc]ông an cấp trên(?:\s+chủ quản)?", "Công an tỉnh Đắk Lắk"),
        (r"CÔNG AN CẤP TRÊN(?:\s+CHỦ QUẢN)?", "Công an tỉnh Đắk Lắk"),
        (r"CÔNG AN TỈNH ĐẮK LẮK", "Công an tỉnh Đắk Lắk"),
        (r"[Cc]ơ quan chủ quản", "Công an tỉnh Đắk Lắk"),
        (r"[Cc]ông an viên", "Cán bộ"),
        (r"CÔNG AN VIÊN", "CÁN BỘ"),
        (r"công an viên", "cán bộ"),
        (r"\s*\[\d+\]", ""),
        (r"(?i)\[CẦN[\s_]BỔ[\s_]SUNG(?::[^\]]*)?\]", "......"),
    ]

    for pattern, replacement in replacements:
        text = re.sub(pattern, replacement, text)

    # 3. FORCE 1.27cm indent + justify on <p> tags OUTSIDE tables
    table_blocks = []
    TABLE_PH = "___TABLE_BLOCK___"

    def stash_table(m):
        block = m.group(0)
        if len(table_blocks) == 0:
            block = re.sub(r"Công an tỉnh Đắk Lắk", "CÔNG AN TỈNH ĐẮK LẮK", block, flags=re.IGNORECASE)
        table_blocks.append(block)
        return TABLE_PH

    text = re.sub(r"<table[\s\S]*?</table>", stash_table, text, flags=re.IGNORECASE)

    def add_paragraph_styles(m):
        tag_attrs = m.group(1) or ""
        content = m.group(2)
        
        style_match = re.search(r"style=['\"]([^'\"]*)['\"]", tag_attrs)
        style_content = style_match.group(1) if style_match else ""
        
        # Check if the paragraph is centered
        is_centered = "text-align: center" in style_content or "text-align:center" in style_content
        
        # Check if paragraph content starts with Roman numerals (e.g. I., II., III.) or A-E followed by a dot
        clean_text = re.sub(r"<[^>]+>", "", content).strip()
        is_section_heading = bool(re.match(r"^(?:[I|V|X]+\.|[A-E]\.)\s", clean_text))
        
        if is_section_heading:
            is_centered = False
            # Clean centering from style and align left/justify with 1.27cm indent
            style_content = re.sub(r"text-align:\s*center;?", "", style_content)
            if "text-indent" not in style_content:
                style_content = "text-indent:1.27cm;text-align:justify;" + style_content
        else:
            if not is_centered:
                if "text-indent" not in style_content:
                    style_content = "text-indent:1.27cm;text-align:justify;" + style_content
                elif "text-align" not in style_content:
                    style_content = "text-align:justify;" + style_content
                    
        style_content = style_content.strip()
        if style_content:
            tag_attrs_clean = re.sub(r"style=['\"]([^'\"]*)['\"]", "", tag_attrs).strip()
            new_tag = f"<p {tag_attrs_clean} style=\"{style_content}\"".replace("  ", " ").strip()
        else:
            new_tag = "<p"
            if tag_attrs.strip():
                new_tag += " " + tag_attrs.strip()
                
        if not new_tag.endswith(">"):
            new_tag += ">"
        return f"{new_tag}{content}</p>"

    text = re.sub(r"<p(\s[^>]*)?>([\s\S]*?)</p>", add_paragraph_styles, text)

    # Restore table blocks (untouched)
    for block in table_blocks:
        text = text.replace(TABLE_PH, block, 1)

    return text

app = FastAPI(
    title="Decree 30 Document Generator API",
    description="API sinh văn bản hành chính theo Nghị định 30/2020/NĐ-CP, sử dụng Google Gemini API.",
    version="4.0.0",
)

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ─── Session State ─────────────────────────────────────────────────────────────

class SessionState:
    """In-memory session state for active sources and conversation."""
    def __init__(self):
        self.sources: List[dict] = []
        self.chat = None

    def add_source(self, source_id: str, source_type: str, name: str, preview: str = "", file_name: str = None, content: str = None):
        self.sources.append({
            "id": source_id,
            "type": source_type,
            "name": name,
            "preview": preview[:200] if preview else "",
            "file_name": file_name, # The genai.File.name
            "content": content
        })

    def reset(self):
        self.sources = []
        self.chat = None


session = SessionState()


# ─── Models ────────────────────────────────────────────────────────────────────

class UploadResponse(BaseModel):
    message: str
    filename: str
    notebook_id: str
    source_id: str

class AddTextRequest(BaseModel):
    text: str
    title: Optional[str] = None

class AddTextResponse(BaseModel):
    message: str
    source_id: str
    notebook_id: str
    title: str

class AddUrlRequest(BaseModel):
    url: str

class AddUrlResponse(BaseModel):
    message: str
    source_id: str
    notebook_id: str
    url: str

class SourceInfo(BaseModel):
    id: str
    type: str
    name: str
    preview: str

class SourceListResponse(BaseModel):
    notebook_id: Optional[str]
    sources: List[SourceInfo]

class GenerateRequest(BaseModel):
    prompt: str
    document_type: str
    notebook_id: Optional[str] = None

class GenerateResponse(BaseModel):
    generated_text: str
    notebook_id: str
    conversation_id: Optional[str] = None

class RefineRequest(BaseModel):
    instruction: str
    current_text: str
    document_type: str
    notebook_id: Optional[str] = None
    conversation_id: Optional[str] = None

class RefineResponse(BaseModel):
    refined_text: str
    notebook_id: str
    conversation_id: Optional[str] = None

class ResetResponse(BaseModel):
    message: str


# ─── Health Check ──────────────────────────────────────────────────────────────

@app.get("/health")
async def health_check():
    return {
        "status": "ok",
        "service": "Decree30 Gemini API",
        "sources_count": len(session.sources),
        "api_key_configured": bool(os.environ.get("GEMINI_API_KEY"))
    }


# ─── Upload File Endpoint ─────────────────────────────────────────────────────

@app.post("/api/upload", response_model=UploadResponse)
async def upload_file(file: UploadFile = File(...)):
    """
    Upload a reference file. Uses Gemini File API for PDF/TXT.
    Extracts text locally for DOCX.
    """
    if not os.environ.get("GEMINI_API_KEY"):
        raise HTTPException(status_code=500, detail="Chưa cấu hình GEMINI_API_KEY. Vui lòng thêm biến môi trường này.")

    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, file.filename)
    source_id = uuid.uuid4().hex[:8]

    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        if file.filename.lower().endswith(".docx"):
            import docx
            doc = docx.Document(file_path)
            fullText = [para.text for para in doc.paragraphs]
            content = '\n'.join(fullText)
            session.add_source(source_id, "file", file.filename, content[:200], content=content)
        else:
            # For PDF, TXT, CSV, etc.
            print(f"Uploading {file.filename} to Gemini...")
            genai_file = genai.upload_file(file_path, display_name=file.filename)
            session.add_source(source_id, "file", file.filename, "Tài liệu được tải lên Gemini (PDF/TXT)", file_name=genai_file.name)

        return UploadResponse(
            message=f"File '{file.filename}' uploaded successfully",
            filename=file.filename,
            notebook_id="gemini-session",
            source_id=source_id,
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)
        if os.path.exists(temp_dir):
            os.rmdir(temp_dir)


# ─── Add Text Source ───────────────────────────────────────────────────────────

@app.post("/api/add-text", response_model=AddTextResponse)
async def add_text_source(request: AddTextRequest):
    if not request.text.strip():
        raise HTTPException(status_code=400, detail="Text cannot be empty")

    title = request.title or f"Text-{uuid.uuid4().hex[:6]}"
    source_id = uuid.uuid4().hex[:8]
    
    session.add_source(source_id, "text", title, request.text[:200], content=request.text)

    return AddTextResponse(
        message=f"Text source '{title}' added successfully",
        source_id=source_id,
        notebook_id="gemini-session",
        title=title,
    )


# ─── Add URL Source ────────────────────────────────────────────────────────────

@app.post("/api/add-url", response_model=AddUrlResponse)
async def add_url_source(request: AddUrlRequest):
    if not request.url.strip():
        raise HTTPException(status_code=400, detail="URL cannot be empty")

    source_id = uuid.uuid4().hex[:8]

    try:
        async with httpx.AsyncClient() as client:
            resp = await client.get(request.url, timeout=15.0)
            resp.raise_for_status()
            soup = BeautifulSoup(resp.text, 'html.parser')
            content = soup.get_text(separator='\n', strip=True)

        session.add_source(source_id, "url", request.url, content[:200], content=content)

        return AddUrlResponse(
            message="URL source added successfully",
            source_id=source_id,
            notebook_id="gemini-session",
            url=request.url,
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Không thể đọc URL: {e}")


# ─── List Sources ─────────────────────────────────────────────────────────────

@app.get("/api/sources", response_model=SourceListResponse)
async def list_sources():
    return SourceListResponse(
        notebook_id="gemini-session",
        sources=[SourceInfo(**s) for s in session.sources],
    )


# ─── Generate Document ────────────────────────────────────────────────────────

@app.post("/api/generate", response_model=GenerateResponse)
async def generate_document(request: GenerateRequest):
    if not os.environ.get("GEMINI_API_KEY"):
        raise HTTPException(status_code=500, detail="Chưa cấu hình GEMINI_API_KEY.")

    try:
        model = genai.GenerativeModel("gemini-2.5-flash") # Or gemini-2.5-pro
        
        contents = []
        # Append all sources
        for src in session.sources:
            if src.get("file_name"):
                # Fetch the file object by name
                f = genai.get_file(src["file_name"])
                contents.append(f)
            if src.get("content"):
                contents.append(f"--- NGUỒN TÀI LIỆU: {src['name']} ---\n{src['content']}")

        # Build prompt
        generation_prompt = build_generation_prompt(
            user_prompt=request.prompt,
            document_type=request.document_type,
        )
        contents.append(generation_prompt)

        print("Sending generation request to Gemini...")
        session.chat = model.start_chat(history=[])
        response = session.chat.send_message(contents)

        generated_text = clean_ai_html_output(response.text)

        return GenerateResponse(
            generated_text=generated_text,
            notebook_id="gemini-session",
            conversation_id="gemini-chat",
        )

    except Exception as e:
        print(f"Error generating: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ─── Refine Document (Chat-based) ─────────────────────────────────────────────

@app.post("/api/refine", response_model=RefineResponse)
async def refine_document(request: RefineRequest):
    if not session.chat:
        raise HTTPException(status_code=400, detail="Chưa có phiên làm việc nào. Hãy tạo văn bản trước.")

    try:
        refinement_prompt = build_refinement_prompt(
            instruction=request.instruction,
            current_text=request.current_text,
            document_type=request.document_type,
        )

        print(f"Refining document with Gemini...")
        response = session.chat.send_message(refinement_prompt)

        refined_text = clean_ai_html_output(response.text)

        return RefineResponse(
            refined_text=refined_text,
            notebook_id="gemini-session",
            conversation_id="gemini-chat",
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ─── Reset Session ─────────────────────────────────────────────────────────────

@app.post("/api/reset", response_model=ResetResponse)
async def reset_session():
    # Cleanup files on Gemini servers if possible
    try:
        for src in session.sources:
            if src.get("file_name"):
                genai.delete_file(src["file_name"])
    except:
        pass
    
    session.reset()
    return ResetResponse(message="Session reset successfully")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
